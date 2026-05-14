from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ========== 全局样式 ==========
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ========== 页边距 ==========
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ========== 辅助函数 ==========
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.left_indent = Cm(1)
    return p

def add_bullet(text, level=0, bold=False):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    for run in p.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if bold:
            run.bold = True
    return p

def add_number(text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    for run in p.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠ {text}")
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 100, 0)
    run.italic = True
    return p

def add_step(num, title, content_lines, url=None, web_desc=None):
    add_heading_styled(f"第 {num} 步：{title}", level=2)
    for line in content_lines:
        if isinstance(line, tuple):
            cmd, desc = line
            add_bullet(desc)
            add_code(cmd)
        else:
            add_para(line)
    if url:
        add_para(f"网站：{url}", bold=True, color=(0, 90, 180))

def add_divider():
    p = doc.add_paragraph()
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.font.size = Pt(8)

# =============================================
# 封面
# =============================================
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("bq-star.com 网站部署与域名绑定指南")
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0, 50, 120)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("使用 Vercel 部署 + Cloudflare DNS/SSL 方案")
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

for _ in range(3):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("项目：mycompanysite（Astro + Vercel）\n日期：2025 年 5 月")
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# =============================================
# 目录页
# =============================================
add_heading_styled("目录", level=1)
toc_items = [
    "概述与架构说明",
    "第 1 步：将代码推送到 GitHub",
    "第 2 步：在 Vercel 部署网站",
    "第 3 步：在 Vercel 绑定自定义域名",
    "第 4 步：在 Cloudflare 添加站点并配置 DNS",
    "第 5 步：更新域名 NS 记录",
    "第 6 步：验证与 SSL/TLS 设置",
    "第 7 步：（可选）性能优化设置",
]
for item in toc_items:
    add_para(item, size=12)

doc.add_page_break()

# =============================================
# 概述
# =============================================
add_heading_styled("概述与架构说明", level=1)

add_para("本指南将指导你完成以下完整流程：将本地 Astro 网站项目部署到 Vercel，然后将自己的域名 bq-star.com 绑定到该网站，并通过 Cloudflare 提供 DNS 解析、CDN 加速和 SSL 证书。")

add_para("最终架构示意图：", bold=True)
add_para("用户 → bq-star.com → Cloudflare DNS → Vercel 托管 → Astro 网站")
add_para("Cloudflare 负责：DNS 解析、DDoS 防护、SSL 证书、CDN 缓存加速")
add_para("Vercel 负责：Astro 应用托管、Serverless 函数运行、自动部署")

doc.add_page_break()

# =============================================
# 第 1 步
# =============================================
add_heading_styled("第一步：将代码推送到 GitHub", level=1)

add_para(
    "在部署到 Vercel 之前，需要先将本地代码推送到 GitHub 仓库。"
    "如果你已经有一个 GitHub 仓库并且代码已在其中，可以跳过此步。"
)

add_heading_styled("1.1 在 GitHub 创建新仓库", level=2)
add_para("网站：https://github.com/new", bold=True, color=(0, 90, 180))

steps_1_1 = [
    "登录你的 GitHub 账号（如果没有，先注册 github.com）",
    "点击右上角 + 号 → New repository",
    "Repository name 输入：mycompanysite",
    "Description（可选）：重庆补全星科技有限公司官网",
    "可见性选择 Private 或 Public（推荐 Public，Vercel 免费计划需要）",
    "不要勾选任何初始化选项（Add a README、.gitignore、license 都不选）",
    "点击 Create repository",
]
for s in steps_1_1:
    add_bullet(s)

add_note("如果之前已经关联了远程仓库，请先检查：git remote -v，如果已有则跳过 1.2")

add_heading_styled("1.2 推送本地代码到 GitHub", level=2)

commands_1_2 = [
    ("git remote add origin https://github.com/你的用户名/mycompanysite.git", "关联远程仓库（将「你的用户名」替换为你的 GitHub 用户名）"),
    ("git branch -M main", "确保分支名为 main"),
    ("git push -u origin main", "推送代码到 GitHub"),
]

for cmd, desc in commands_1_2:
    add_para(desc)
    add_code(cmd)

add_para("完成后，刷新 GitHub 仓库页面，应该能看到所有代码文件。", bold=True)

doc.add_page_break()

# =============================================
# 第 2 步
# =============================================
add_heading_styled("第二步：在 Vercel 部署网站", level=1)

add_para("网站：https://vercel.com", bold=True, color=(0, 90, 180))

add_heading_styled("2.1 注册并登录 Vercel", level=2)
steps_2_1 = [
    "打开 vercel.com",
    "点击右上角 Sign Up（或 Login，如果你已有账号）",
    "推荐使用 Continue with GitHub 一键登录，这样可以自动导入你的 GitHub 仓库",
    "授权 Vercel 访问你的 GitHub 仓库",
]
for s in steps_2_1:
    add_bullet(s)

add_heading_styled("2.2 导入并部署项目", level=2)
steps_2_2 = [
    "登录后进入 Dashboard，点击 Add New → Project",
    "在 Import Git Repository 列表中找到 mycompanysite，点击 Import",
    "如果列表中没有，点击 Configure GitHub App 或 Adjust GitHub App Permissions，确保仓库可见",
    "Framework Preset：Vercel 会自动检测为 Astro，无需修改",
    "Root Directory：保持默认（项目根目录）",
    "Build and Output Settings：无需修改（Astro 的 vercel adapter 会自动处理）",
    "Environment Variables：无需添加（本项目没有需要设置的密钥）",
    "点击 Deploy 按钮",
]
for s in steps_2_2:
    add_bullet(s)

add_heading_styled("2.3 等待部署完成", level=2)
add_para("部署过程大约需要 1-3 分钟。你会看到实时构建日志。")
add_para("部署成功后，你会看到以下页面信息：")
steps_2_3 = [
    "Congratulations! 页面",
    "自动分配的临时域名：https://mycompanysite.vercel.app",
    "Production Deployment 标记为 Ready",
    "点击 Visit 按钮可以预览网站",
]
for s in steps_2_3:
    add_bullet(s)

add_note("请记下这个 Vercel 分配的域名（xxx.vercel.app），下一步会用到。")

doc.add_page_break()

# =============================================
# 第 3 步
# =============================================
add_heading_styled("第三步：在 Vercel 绑定自定义域名", level=1)

add_para("网站：https://vercel.com/dashboard", bold=True, color=(0, 90, 180))
add_para("进入你的项目 → Settings → Domains")

add_heading_styled("3.1 添加域名", level=2)
steps_3_1 = [
    "在 Vercel Dashboard 中点击你的 mycompanysite 项目",
    "进入顶部导航的 Settings 标签",
    "左侧菜单选择 Domains",
    "输入框中输入 bq-star.com，点击 Add",
    "Vercel 会提示 DNS 配置方式，选择：Add DNS record in Cloudflare later",
]
for s in steps_3_1:
    add_bullet(s)

add_heading_styled("3.2 记录 Vercel 提供的 DNS 目标", level=2)
add_para("添加域名后，Vercel 会在 Domains 列表中显示一条记录，类似：", bold=True)
add_code("bq-star.com  →  cname.vercel-dns.com")
add_para("或者有时是：", bold=True)
add_code("bq-star.com  →  mycompanysite.vercel.app")
add_para("请记录这个目标地址，下一步配置 Cloudflare DNS 时需要用到。", bold=True)
add_note("如果 Vercel 要求你配置 Nameserver，先忽略——我们会在 Cloudflare 统一管理 DNS。")

doc.add_page_break()

# =============================================
# 第 4 步
# =============================================
add_heading_styled("第四步：在 Cloudflare 添加站点并配置 DNS", level=1)

add_para("网站：https://dash.cloudflare.com", bold=True, color=(0, 90, 180))

add_heading_styled("4.1 注册 Cloudflare", level=2)
steps_4_1 = [
    "打开 dash.cloudflare.com",
    "点击 Sign Up，输入邮箱地址并设置密码",
    "验证邮箱后登录",
]
for s in steps_4_1:
    add_bullet(s)

add_heading_styled("4.2 添加站点", level=2)
steps_4_2 = [
    "登录后点击 Add a Site（蓝色按钮）",
    "输入你的域名：bq-star.com",
    "点击 Add Site",
    "选择 Free 计划（免费计划已包含 DNS 解析、CDN、SSL 证书）",
    "点击 Continue",
]
for s in steps_4_2:
    add_bullet(s)

add_heading_styled("4.3 配置 DNS 记录", level=2)
add_para("Cloudflare 会自动扫描现有的 DNS 记录（由于是新域名，可能没有或很少）。")
add_para("现在手动添加以下两条记录：", bold=True)

add_para("记录 1：主域名", bold=True, size=12)
table = doc.add_table(rows=5, cols=2, style='Light Shading Accent 1')
table.cell(0,0).text = "类型"
table.cell(0,1).text = "CNAME"
table.cell(1,0).text = "名称"
table.cell(1,1).text = "@"
table.cell(2,0).text = "目标"
table.cell(2,1).text = "cname.vercel-dns.com（或 Vercel 提供的目标地址）"
table.cell(3,0).text = "代理状态"
table.cell(3,1).text = "开启（橙色云图标）"
table.cell(4,0).text = "TTL"
table.cell(4,1).text = "Auto"

doc.add_paragraph()

add_para("记录 2：www 子域名", bold=True, size=12)
table2 = doc.add_table(rows=5, cols=2, style='Light Shading Accent 1')
table2.cell(0,0).text = "类型"
table2.cell(0,1).text = "CNAME"
table2.cell(1,0).text = "名称"
table2.cell(1,1).text = "www"
table2.cell(2,0).text = "目标"
table2.cell(2,1).text = "bq-star.com"
table2.cell(3,0).text = "代理状态"
table2.cell(3,1).text = "开启（橙色云图标）"
table2.cell(4,0).text = "TTL"
table2.cell(4,1).text = "Auto"

doc.add_paragraph()
add_para("操作提示：", bold=True)
add_bullet("点击 Add Record 添加新记录")
add_bullet("类型选择 CNAME")
add_bullet("名称输入 @（代表根域名）或 www")
add_bullet("目标输入 Vercel 提供的地址")
add_bullet("确保代理开关是橙色云状态（表示开启 Cloudflare 代理）")
add_bullet("点击 Save")

add_note("如果 Vercel 提供的是 A 记录 IP 地址而非 CNAME 目标，则类型选择 A 记录，值填入对应的 IP。以 Vercel Domains 页面显示为准。")

doc.add_page_break()

# =============================================
# 第 5 步
# =============================================
add_heading_styled("第五步：更新域名 NS 记录（最关键的一步）", level=1)

add_para("这一步让 bq-star.com 的 DNS 管理权从当前注册商转移到 Cloudflare。", bold=True)

add_heading_styled("5.1 获取 Cloudflare 的 Nameserver 地址", level=2)
add_para("在 Cloudflare 完成添加站点和 DNS 配置后，它会自动分配两个 Nameserver（名称服务器）地址，类似：")
add_code("dara.ns.cloudflare.com")
add_code("kanye.ns.cloudflare.com")
add_para("每个站点分配的两个地址不同，以你实际看到的为准。")

add_heading_styled("5.2 到域名注册商处修改 NS 记录", level=2)
add_para("你需要到你购买 bq-star.com 的注册商网站去操作。")
add_para("常见注册商及其操作位置：", bold=True)

registrars = [
    "阿里云（万网）：登录 aliyun.com → 控制台 → 域名 → 域名列表 → 管理 → DNS 修改",
    "腾讯云：登录 console.cloud.tencent.com → 域名管理 → DNS 修改",
    "Namecheap：登录 → Domain List → Manage → Nameservers → Custom DNS",
    "GoDaddy：登录 → 我的域名 → DNS → Nameservers → 自定义",
    "Namesilo：登录 → Domain Manager → 对应域名 → Change NS",
]
for r in registrars:
    add_bullet(r)

add_heading_styled("5.3 替换 Nameserver", level=2)
steps_5_3 = [
    "在注册商的 DNS 管理页面，将原有的 NS 记录替换为 Cloudflare 提供的那两个地址",
    "删除原有的所有 NS 记录（通常是注册商默认的 NS）",
    "只保留 Cloudflare 的两个 NS 地址",
    "保存设置",
]
for s in steps_5_3:
    add_bullet(s)

add_heading_styled("5.4 等待生效", level=2)
add_para("NS 记录变更的生效时间：")
add_bullet("通常 5-30 分钟内开始生效")
add_bullet("最长可能需要 48 小时（全球 DNS 缓存刷新）")
add_bullet("大多数情况下 1 小时内即可完成")

add_para("在 Cloudflare 仪表盘中，你会看到站点状态从 Pending 变为 Active。", bold=True)
add_note("在生效期间，网站仍然可以正常工作——DNS 解析会逐步切换到 Cloudflare，不会出现停机。")

doc.add_page_break()

# =============================================
# 第 6 步
# =============================================
add_heading_styled("第六步：验证与 SSL/TLS 设置", level=1)

add_heading_styled("6.1 确认部署和域名解析正常", level=2)
add_para("等待几分钟后，用以下方法测试：")
add_bullet("在浏览器访问 https://bq-star.com 是否能正常打开网站")
add_bullet("访问 https://www.bq-star.com 是否能正常跳转或打开")
add_bullet("也可以使用命令行测试：", bold=True)
add_code("curl -I https://bq-star.com")

add_heading_styled("6.2 在 Vercel 确认域名状态", level=2)
add_para("回到 Vercel → Project → Settings → Domains，查看 bq-star.com 的状态：")
add_bullet("状态显示 Valid，旁边有绿色勾 ✅")
add_bullet("如果显示 Pending 或 Configuration，说明 DNS 还在传播中，等待即可")
add_bullet("Vercel 会自动检查 DNS 配置并签发 SSL 证书")

add_heading_styled("6.3 Cloudflare SSL/TLS 设置", level=2)
add_para("在 Cloudflare 仪表盘中，进入 SSL/TLS → Overview：")

steps_6_3 = [
    "SSL/TLS encryption mode：选择 Full (strict) 模式",
    "Full (strict) 表示 Cloudflare 到 Vercel 之间也是加密连接，且会验证证书有效性",
    "这是最安全的模式，推荐使用",
    "不要选择 Flexible 模式（该模式 Cloudflare 到服务器是明文传输）",
]
for s in steps_6_3:
    add_bullet(s)

add_para("其他推荐开启的 SSL 选项：", bold=True)
add_bullet("Always Use HTTPS：强制 HTTP 请求跳转到 HTTPS")
add_bullet("Automatic HTTPS Rewrites：自动将页面中的 HTTP 链接改写为 HTTPS")
add_bullet("以上两个选项都在 SSL/TLS → Edge Certificates 页面中")

add_para("Cloudflare 会自动为你的域名签发并管理 SSL 证书，无需手动申请或续期。", bold=True)

doc.add_page_break()

# =============================================
# 第 7 步（可选）
# =============================================
add_heading_styled("第七步：（可选）性能优化设置", level=1)

add_para("以下为推荐在 Cloudflare 中开启的优化选项：")

optimizations = [
    ("Speed → Optimization", ["Auto Minify：开启 JavaScript、CSS、HTML 压缩",
                               "Brotli：开启（默认已开启，与 astro-compressor 的 brotli 压缩兼容）",
                               "Polish：开启 Lossless 模式（自动优化图片）"]),
    ("Speed → Speed", ["开启 Early Hints", "开启 0-RTT Connection Resumption"]),
    ("Security → Settings", ["Security Level 设为 Medium",
                              "Bot Fight Mode 按需开启（免费计划可用）"]),
]

for section, items in optimizations:
    add_para(section, bold=True, size=12)
    for item in items:
        add_bullet(item)

doc.add_paragraph()
add_para("完成以上所有步骤后，你的 bq-star.com 网站将正式上线，", bold=True)
add_para("通过 Cloudflare 提供 CDN 加速 + DDoS 防护 + SSL 加密，", bold=True)
add_para("网站应用由 Vercel 持续部署和托管。", bold=True)

doc.add_paragraph()
add_para("至此，部署与绑定工作全部完成。", bold=True, size=14, color=(0, 120, 0))

# ========== 保存 ==========
output_path = os.path.expanduser("~/Desktop/bq-star-com-deploy-guide.docx")
# Try to save to a reasonable location
try:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")
except:
    output_path = "/data/work/web/mycompanysite/bq-star-com-deploy-guide.docx"
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")
